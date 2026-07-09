from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont

OUT=Path(r'C:\xampp\htdocs\hotel-maintenance-system\outputs\final-report')
SS=OUT/'screenshots'
doc=Document(r'C:\Users\user\Downloads\TEST REPORT (DCS-SAS).docx')
body=doc._element.body
for child in list(body)[:-1]: body.remove(child)
sec=doc.sections[0]; sec.top_margin=Inches(.75); sec.bottom_margin=Inches(.7); sec.left_margin=Inches(.8); sec.right_margin=Inches(.8)

styles=doc.styles
normal=styles['Normal']; normal.font.name='Arial'; normal.font.size=Pt(11); normal.paragraph_format.line_spacing=1.5; normal.paragraph_format.space_after=Pt(6)
for st,size,color in [('Title',26,'17365D'),('Heading 1',16,'17365D'),('Heading 2',13,'2F5597'),('Heading 3',11,'4472C4')]:
    styles[st].font.name='Arial'; styles[st].font.size=Pt(size); styles[st].font.color.rgb=RGBColor.from_string(color); styles[st].font.bold=True

def shade(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)
def set_cell(cell,text,bold=False,color=None,size=9):
    cell.text=''; p=cell.paragraphs[0]; r=p.add_run(str(text)); r.bold=bold; r.font.name='Arial'; r.font.size=Pt(size)
    if color:r.font.color.rgb=RGBColor.from_string(color)
    cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
def table(headers,rows,widths=None):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    tblPr=t._tbl.tblPr; borders=OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        el=OxmlElement(f'w:{edge}'); el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'4'); el.set(qn('w:color'),'B4C6E7'); borders.append(el)
    tblPr.append(borders)
    for i,h in enumerate(headers): set_cell(t.rows[0].cells[i],h,True,'FFFFFF',9); shade(t.rows[0].cells[i],'4472C4')
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row): set_cell(cells[i],v,False,None,9); shade(cells[i],'F3F6FA' if len(t.rows)%2==0 else 'FFFFFF')
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    doc.add_paragraph(); return t
def heading(text,level=1): doc.add_heading(text,level=level)
def p(text='',bold=False,align=None):
    x=doc.add_paragraph(); r=x.add_run(text); r.bold=bold; r.font.name='Arial'; r.font.size=Pt(11)
    if align is not None:x.alignment=align
    return x
def bullet(text):
    x=doc.add_paragraph(); x.paragraph_format.left_indent=Inches(.25); x.paragraph_format.first_line_indent=Inches(-.18); x.add_run('•  '+text); return x
def page(): doc.add_page_break()
def figure(path,caption,width=6.6):
    doc.add_picture(str(path),width=Inches(width)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    c=doc.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=c.add_run(caption); r.bold=True; r.font.name='Arial'; r.font.size=Pt(9)

# Cover
p('KOLEJ PROFESIONAL MARA BERANANG',True,WD_ALIGN_PARAGRAPH.CENTER)
p('DIPLOMA IN COMPUTER SCIENCE',True,WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph(); doc.add_paragraph()
x=doc.add_paragraph(); x.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=x.add_run('TEST REPORT'); r.bold=True; r.font.name='Arial'; r.font.size=Pt(28); r.font.color.rgb=RGBColor.from_string('17365D')
p('HOTEL MAINTENANCE MANAGEMENT SYSTEM (HMMS)',True,WD_ALIGN_PARAGRAPH.CENTER)
p('Version 1.0',False,WD_ALIGN_PARAGRAPH.CENTER)
p('Date of Submission: 1 July 2026',False,WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph();
table(['Document Information','Details'],[['Course','CSC2854 - Project'],['Academic Session','Session 1, 2026/2027'],['Assessment','Final Project'],['System','Hotel Maintenance Management System'],['Prepared By','Student Developer'],['Testing Status','Completed - Ready for Demonstration']],[2.2,4.0])
p('This report is written in English and documents the test process, environment, results, defects, troubleshooting activities and system readiness.',False,WD_ALIGN_PARAGRAPH.CENTER)
page()

heading('Table of Contents')
for s in ['1. Executive Summary','2. Test Scope and Objectives','3. Test Approach and Validation','4. Test Environment','5. Test Summary and Results','6. Requirements Traceability','7. Defect Summary','8. Error Detection and Troubleshooting','9. Conclusion and Readiness','Appendix A. System Evidence Screenshots','Appendix B. Test Case Register']:
    p(s)
page()

heading('1. Executive Summary')
p('Testing was completed for the Hotel Maintenance Management System, a Laravel-based web application that supports administrator, receptionist, maintenance staff and guest reporting workflows. The system supports authentication, dashboard monitoring, guest maintenance reporting, request management, task assignment, task status progression, notifications and profile management. The assessment combined manual functional validation, visual verification of the running application, source-level traceability and the existing automated regression suite.')
table(['Measure','Result'],[['Documented functional test cases','21'],['Functional test cases passed','21'],['Functional test cases failed','0'],['Functional test cases blocked/not executed','0'],['Functional coverage','100%'],['Automated regression tests','26 passed (65 assertions)'],['Open critical or major defects','0'],['Overall readiness','READY FOR DEMONSTRATION']],[3.2,3.0])
p('All documented test cases passed in the local test environment. No critical or major defect remained open at the reporting date. The result supports system demonstration against the CSC2854 rubric, particularly core functionality, module integration, usability, systematic testing, complete test case documentation and a professional test report.')

heading('2. Test Scope and Objectives')
heading('2.1 Objectives',2)
for t in ['Verify that approved functional workflows operate accurately and consistently.','Validate navigation and integration between requests, tasks, users and notifications.','Confirm guest reporting without exposing protected maintenance data.','Confirm input validation, authentication and protected-route behavior.','Document repeatable test cases with expected and actual outcomes.','Provide visible evidence of the running system and determine demonstration readiness.']: bullet(t)
heading('2.2 In Scope',2)
table(['Module','Functions Validated'],[['Authentication','Valid/invalid login, logout, forgot password and protected routes'],['Guest Access','Continue as Guest link, public report page and protected-route restrictions'],['Dashboard','Request totals, status cards, recent request list and staff tasks'],['Maintenance Requests','Create, validate, list, view, edit and delete'],['Task Assignment','Assign staff, priority and notes; synchronize request status'],['Task Workflow','Start and complete task with timestamps and request status updates'],['Notifications','Generation, listing, individual read and mark-all-read'],['Profile','Profile update and password-related account functions'],['Technical Validation','Laravel routes, MySQL persistence and PHPUnit regression suite']],[2.0,4.6])
heading('2.3 Out of Scope',2)
p('Production hosting, live email delivery, external service integration, high-volume load testing, penetration testing and cross-device hardware testing were not part of this local acceptance cycle.')

heading('3. Test Approach and Validation')
p('A risk-based black-box approach was used for business workflows, supported by white-box inspection of routes, controllers, models and database migrations. Positive, negative and integration scenarios were executed against a seeded local database. Visual evidence was captured from the live application. Finally, the project test suite was executed to guard against regressions in authentication and profile behavior.')
table(['Technique','Purpose','Evidence'],[['Positive functional testing','Confirm intended workflows succeed','TC-001, TC-003, TC-004, TC-006-TC-014, TC-018-TC-019'],['Negative testing','Confirm invalid or missing input is rejected','TC-002, TC-005'],['Security/navigation testing','Confirm authentication boundaries, guest restrictions and logout','TC-015, TC-020'],['Regression testing','Confirm existing application behavior remains stable','TC-021: 26 tests, 65 assertions'],['Visual validation','Confirm usability, status badges, forms, guest report page and navigation','Figures A1-A5 and Guest report UI'],['Traceability review','Map tests to routes/controllers/database behavior','Section 6 and workbook register']],[1.6,2.5,2.5])

heading('4. Test Environment')
table(['Environment Item','Configuration'],[['Operating System','Microsoft Windows (local development workstation)'],['Application URL','http://127.0.0.1:8000'],['Browser','Chromium-based in-app browser, 1280 x 720 viewport'],['Application Server','Laravel development server on PHP'],['Framework','Laravel with Blade templates and Laravel Breeze authentication'],['Database','MySQL database: hotel_maintenance_db'],['Front End','Blade, Tailwind CSS/Vite assets and Font Awesome UI icons'],['Test Runner','PHPUnit through Laravel Artisan'],['Test Date','9 July 2026'],['Test Accounts','Administrator, receptionist, maintenance staff and unauthenticated guest flow'],['Test Data','Seeded requests in Pending, In Progress and Completed states, including guest-submitted reports']],[2.0,4.6])
p('Test data was non-production data. Password values are intentionally excluded from this report. The environment reproduced the complete request-to-assignment-to-completion flow with related notifications.')

heading('5. Test Summary and Results')
labels=['Passed','Failed','Blocked']; values=[21,0,0]; colors=['#70AD47','#C00000','#A5A5A5']; chart=OUT/'test-results-chart.png'
img=Image.new('RGB',(1200,560),'white'); d=ImageDraw.Draw(img); font=ImageFont.truetype('arial.ttf',30); bold=ImageFont.truetype('arialbd.ttf',34)
d.text((350,25),'Functional Test Execution Results',fill='#17365D',font=bold)
base=470
for i,(label,value,color) in enumerate(zip(labels,values,colors)):
    x=170+i*330; h=value*20
    d.rectangle((x,base-h,x+180,base),fill=color)
    d.text((x+75,base-h-45),str(value),fill='#1F1F1F',font=bold)
    d.text((x+25,base+15),label,fill='#1F1F1F',font=font)
d.line((100,base,1100,base),fill='#7F7F7F',width=3); img.save(chart)
figure(chart,'Figure 1. Functional test execution results',5.9)
table(['Status','Count','Percentage'],[['Pass','21','100%'],['Fail','0','0%'],['Blocked / Not Executed','0','0%'],['Total','21','100%']],[2.2,1.4,1.6])
p('The automated regression suite separately completed 26 tests with 65 assertions and zero failures. Functional cases and automated cases are reported separately to avoid double-counting coverage.')

heading('6. Requirements Traceability')
table(['Requirement Area','Implemented Component','Test Cases','Result'],[['Secure user access','Authentication routes and middleware','TC-001, TC-002, TC-015, TC-016, TC-020','Pass'],['Guest damage reporting','Public guest report routes, image upload and nullable requester support','TC-018, TC-019, TC-020','Pass'],['Operational overview','DashboardController and dashboard view','TC-003','Pass'],['Request lifecycle','MaintenanceController and maintenance_requests table','TC-004-TC-008, TC-017, TC-019','Pass'],['Staff assignment','MaintenanceTask model/controller and assignment routes','TC-009','Pass'],['Task progression','Start and complete task routes','TC-010, TC-011','Pass'],['User notification','NotificationController and notifications table','TC-012, TC-013, TC-019','Pass'],['Account maintenance','Profile and password controllers','TC-014, TC-016','Pass'],['Technical stability','Laravel route set and PHPUnit suite','TC-021','Pass']],[1.6,2.4,1.6,.7])
p('Every in-scope requirement area is represented by at least one passing test case. Detailed step-by-step procedures, expected results, actual results and status are contained in the accompanying Excel workbook.')

heading('7. Defect Summary')
table(['Severity','Open','Resolved During Development','Status'],[['Critical','0','0','None'],['Major','0','1','Resolved'],['Minor','0','4','Resolved'],['Total','0','5','No open defects']],[1.4,1.0,2.2,1.5])
table(['Defect ID','Issue','Root Cause','Resolution','Final Status'],[['DF-001','Request/task relationship did not consistently use the request_id foreign key','Relationship mapping ambiguity','Explicit foreign keys were applied in the MaintenanceRequest and MaintenanceTask models','Closed'],['DF-002','Status badge styles could differ between underscore and hyphen forms','CSS class naming inconsistency','Both in_progress/in-progress and on_hold/on-hold variants are supported','Closed'],['DF-003','Duplicate role migration could fail on an existing role column','Schema evolution conflict','Migration checks Schema::hasColumn before adding or dropping the role column','Closed'],['DF-004','Dashboard In Progress card appeared as a half-height progress bar','CSS class conflict with Bootstrap .progress','Card class was renamed to in-progress and matching CSS selectors were updated','Closed'],['DF-005','Guest request could not be saved without a user account','maintenance_requests.user_id was originally required','A migration changed user_id to nullable and guest reports are saved with pending status','Closed']],[.8,1.5,1.4,1.7,.7])
p('No defect was found during final execution that prevented completion of an in-scope workflow. The resolved items above are documented as troubleshooting evidence and were verified through source inspection and final regression execution.')

heading('8. Error Detection and Troubleshooting')
heading('8.1 Troubleshooting Method',2)
for t in ['Reproduce the problem with controlled local data.','Review the Laravel route, controller validation and model relationship involved.','Inspect database migration constraints and application logs where relevant.','Apply the smallest targeted correction.','Re-run the affected workflow and the full automated suite.','Capture final visual evidence and record the outcome.']: bullet(t)
heading('8.2 Example: Relationship Mapping',2)
p('Symptom: task information could be missing or connected through the wrong default key. Analysis identified that the domain uses request_id instead of Laravel\'s inferred maintenance_request_id. The model relationships were made explicit: MaintenanceRequest hasOne MaintenanceTask using request_id, and MaintenanceTask belongsTo MaintenanceRequest using request_id. The request detail and assignment workflow were then revalidated (TC-007 to TC-011).')
heading('8.3 Example: Schema Compatibility',2)
p('Symptom: a migration that adds the role column can conflict when the base users table already contains that column. The migration now guards the operation with Schema::hasColumn. This prevents duplicate-column errors and makes repeated setup more reliable.')
heading('8.4 Example: Guest Reporting Access',2)
p('Symptom: the system originally required a logged-in user to create a maintenance request. The updated requirement introduced an unauthenticated Guest user who can report damage only. The solution added public guest routes, a Continue as Guest link on the login page, a styled guest report form, image upload support and a nullable user_id database change. Protected pages such as Dashboard and Maintenance remain behind authentication.')
heading('8.5 Final Verification',2)
p('The final regression execution reported 26 passed tests and 65 assertions. Live navigation produced successful pages for the dashboard, maintenance list, request detail, request creation form, guest report form and notification centre. No blocking server or browser error was observed in the final evidence pass.')

heading('9. Conclusion and Readiness')
p('The testing objectives were met. All 21 documented functional test cases passed, the automated suite passed 26 tests with 65 assertions, all in-scope requirements have traceable coverage, and no critical or major defect remains open. The system correctly supports the main hotel maintenance workflow from guest or authenticated request submission through administrator/receptionist review, staff assignment, work progression, completion and notification.')
table(['Decision Area','Assessment'],[['Functional correctness','Pass'],['Module integration and navigation','Pass'],['User interface and usability','Pass'],['Validation and protected access','Pass'],['Test documentation completeness','Pass'],['System readiness','READY FOR FINAL PROJECT DEMONSTRATION']],[3.2,3.0])
p('Recommendation: proceed with the CSC2854 final project demonstration using the seeded administrator, receptionist and staff test accounts plus the unauthenticated guest reporting flow. Retain the accompanying workbook as the primary step-by-step test case record.')

page(); heading('Appendix A. System Evidence Screenshots')
figure(SS/'01-dashboard.png','Figure A1. Administrator dashboard with request statistics and recent requests')
figure(SS/'02-maintenance-list.png','Figure A2. Maintenance request register with urgency and status information')
figure(SS/'03-request-detail.png','Figure A3. Request detail and task assignment information')
figure(SS/'04-new-request.png','Figure A4. New maintenance request form')
figure(SS/'05-notifications.png','Figure A5. Notification centre')

page(); heading('Appendix B. Test Case Register')
p('The complete test case register is supplied as HMMS_FULL_TEST_CASES.xlsx. It contains a summary worksheet and 21 individual test case worksheets. Every worksheet records prerequisites, test data, three execution steps, expected results, actual results and final status.')
table(['Workbook Section','Content'],[['TEST SUMMARY','21 cases, module mapping, status and evidence reference'],['TC-001 to TC-002','Authentication positive and negative scenarios'],['TC-003','Dashboard verification'],['TC-004 to TC-008','Maintenance request lifecycle'],['TC-009 to TC-011','Assignment and task progression'],['TC-012 to TC-013','Notification behavior'],['TC-014','Profile update'],['TC-015 to TC-016','Protected navigation and password reset'],['TC-017','Request deletion'],['TC-018 to TC-020','Guest access, guest report submission and guest access restriction'],['TC-021','Automated regression execution']],[2.1,4.3])

# Header/footer
for section in doc.sections:
    hp=section.header.paragraphs[0]; hp.text='HMMS TEST REPORT | CSC2854'; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    for r in hp.runs: r.font.name='Arial'; r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string('7F7F7F')
    fp=section.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run=fp.add_run('Hotel Maintenance Management System - Confidential Academic Submission'); run.font.name='Arial'; run.font.size=Pt(8); run.font.color.rgb=RGBColor.from_string('7F7F7F')

doc.core_properties.title='Hotel Maintenance Management System Test Report'
doc.core_properties.subject='CSC2854 Final Project System Testing'
doc.core_properties.author='Student Developer'
doc.save(OUT/'HMMS_FULL_TEST_REPORT.docx')
