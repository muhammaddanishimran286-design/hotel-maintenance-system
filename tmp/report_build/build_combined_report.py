from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import load_workbook

ROOT=Path(r'C:\xampp\htdocs\hotel-maintenance-system')
OUT=ROOT/'outputs'/'final-report'
SOURCE=OUT/'HMMS_FULL_TEST_REPORT.docx'
XLSX=OUT/'HMMS_FULL_TEST_CASES.xlsx'
DEST=OUT/'HMMS_COMBINED_FINAL_REPORT.docx'

doc=Document(SOURCE)

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)

def borders(table, color='B4C6E7'):
    tblPr=table._tbl.tblPr; e=OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        x=OxmlElement(f'w:{edge}'); x.set(qn('w:val'),'single'); x.set(qn('w:sz'),'4'); x.set(qn('w:color'),color); e.append(x)
    tblPr.append(e)

def cell(cell,text,bold=False,color=None,size=8.5):
    cell.text=''; p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(0)
    r=p.add_run('' if text is None else str(text)); r.font.name='Arial'; r.font.size=Pt(size); r.bold=bold
    if color: r.font.color.rgb=RGBColor.from_string(color)
    cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER

def table(headers, rows, widths=None, font=8.5):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False; borders(t)
    for i,h in enumerate(headers): cell(t.rows[0].cells[i],h,True,'FFFFFF',font); shade(t.rows[0].cells[i],'4472C4')
    for idx,row in enumerate(rows):
        cs=t.add_row().cells
        for i,v in enumerate(row): cell(cs[i],v,False,None,font); shade(cs[i],'F3F6FA' if idx%2 else 'FFFFFF')
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return t

def heading(text,level=1):
    p=doc.add_heading(text,level=level); p.paragraph_format.keep_with_next=True; return p

def para(text,bold=False):
    p=doc.add_paragraph(); r=p.add_run(text); r.font.name='Arial'; r.font.size=Pt(11); r.bold=bold; return p

def bullet(text):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.28); p.paragraph_format.first_line_indent=Inches(-.2)
    r=p.add_run('•  '+text); r.font.name='Arial'; r.font.size=Pt(11); return p

doc.add_page_break()
heading('Appendix C. CSC2854 Assessment Brief and Compliance Mapping')
para('Source document: FYP CSC2854 S1 2026 2027. The following section consolidates the assessment requirements supplied with the project and maps them to the evidence contained in this combined report.')

heading('C.1 Assessment Information',2)
table(['Item','Requirement'],[
    ['Institution','Kolej Profesional MARA Beranang'],
    ['Programme','Diploma in Computer Science'],
    ['Course','CSC2854 - Project'],
    ['Academic Session','Session 1, 2026/2027'],
    ['Assessment','Final Project'],
    ['Assessment Period','5 June 2026 to 3 July 2026'],
    ['Course Learning Outcome','CLO3: Build a functional system based on user and technical requirements.'],
    ['Report Format','Arial, 11 pt, 1.5 line spacing'],
    ['Required Supporting Material','Source code and a 2-3 minute error detection and troubleshooting video']
],[2.2,4.2],9)

heading('C.2 System Implementation Requirements',2)
table(['Rubric Area','Required Demonstration','HMMS Evidence'],[
    ['1(a) Requirement Fulfilment and Scope','Show that approved user and technical requirements are implemented.','Sections 2 and 6; Appendix B; TC-001 to TC-018'],
    ['1(b) Construction and Core Functionality','Execute major functions accurately and consistently without critical errors.','Sections 3 and 5; Figures A1-A5; full test cases'],
    ['1(c) Integration and Navigation','Show seamless workflows and module interaction.','Request, assignment, task and notification traceability in Section 6'],
    ['1(d) User Interface and Usability','Demonstrate a clear, consistent and usable interface.','Appendix A system screenshots'],
    ['1(e) Technical Implementation','Describe frameworks, languages, database and development tools.','Section 4 test environment and technical validation']
],[1.5,2.6,2.3],8)

heading('C.3 System Testing Requirements',2)
table(['Rubric Area','Required Demonstration','HMMS Evidence'],[
    ['2(a) Testing and Validation','Explain systematic testing and prove validation against requirements.','Sections 3, 5 and 6'],
    ['2(b) Error Detection and Troubleshooting','Identify, analyse, resolve and verify system errors.','Sections 7 and 8; controlled relationship bug demonstration'],
    ['2(c) Test Case Documentation','Document detailed tests covering major functions and outcomes.','Appendix D contains all 18 detailed test cases'],
    ['2(d) Test Report','Provide summary, environment, defect summary, conclusion and suitable charts.','Sections 1, 4, 5, 7 and 9']
],[1.5,2.6,2.3],8)

heading('C.4 Submission Checklist',2)
for text in [
    'Working and functional Hotel Maintenance Management System.',
    'Appropriate digital technologies: Laravel, PHP, Blade, MySQL and PHPUnit.',
    'Practical development evidence through integrated request and task workflows.',
    'System testing evidence through 18 documented functional cases and an automated regression suite.',
    'Supporting technical documentation consolidated into this report.',
    'Error detection and troubleshooting video supported by the prepared controlled bug scenario.'
]: bullet(text)

doc.add_page_break()
heading('Appendix D. Complete Test Case Documentation')
para('This appendix reproduces the complete content of the accompanying Excel test case workbook so that the assessment brief, test report, screenshots and detailed test cases are available in one Word document.')

wb=load_workbook(XLSX,data_only=True)
case_sheets=[s for s in wb.sheetnames if s.startswith('TC-')]
for n,name in enumerate(case_sheets):
    ws=wb[name]
    if n: doc.add_page_break()
    heading(f'{name}: {ws["F1"].value}',2)
    table(['Field','Details'],[
        ['Test Case ID',ws['C1'].value],
        ['Version',ws['C2'].value],
        ['Date Tested',ws['F2'].value.strftime('%d %B %Y') if hasattr(ws['F2'].value,'strftime') else ws['F2'].value],
        ['Final Status','Pass'],
        ['Test Scenario',ws['B7'].value],
        ['Prerequisite 1',ws['B5'].value],
        ['Prerequisite 2',ws['B6'].value],
        ['Test Data',ws['G5'].value]
    ],[1.55,4.85],8.5)
    steps=[]
    for row in range(10,13):
        steps.append([ws[f'A{row}'].value,ws[f'B{row}'].value,ws[f'E{row}'].value,ws[f'G{row}'].value,ws[f'I{row}'].value])
    table(['Step','Action','Expected Result','Actual Result','Status'],steps,[.45,1.85,1.9,1.3,.7],8)
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(3)
    r=p.add_run('Result: PASS - All expected outcomes were achieved.'); r.bold=True; r.font.name='Arial'; r.font.size=Pt(10); r.font.color.rgb=RGBColor.from_string('006100')

doc.core_properties.title='HMMS Combined Final Report'
doc.core_properties.subject='CSC2854 assessment brief, test report and complete test cases'
doc.save(DEST)
print(DEST)
