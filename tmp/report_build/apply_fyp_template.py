from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT=Path(r'C:\xampp\htdocs\hotel-maintenance-system')
src=ROOT/'outputs'/'final-report'/'HMMS_COMBINED_FINAL_REPORT.docx'
dst=ROOT/'outputs'/'final-report'/'HMMS_FINAL_REPORT_FYP_TEMPLATE.docx'
doc=Document(src)

# Match the CSC2854 assessment front-page wording while retaining the complete report body.
updates={
    0:'KOLEJ PROFESIONAL MARA BERANANG',
    1:'DIPLOMA IN COMPUTER SCIENCE',
    4:'FINAL PROJECT REPORT',
    5:'HOTEL MAINTENANCE MANAGEMENT SYSTEM (HMMS)',
    6:'COURSE NAME: PROJECT     |     COURSE CODE: CSC2854',
    7:'ACADEMIC SESSION: SESSION 1 2026/2027',
    10:'Assessment Period: 5 June 2026 - 3 July 2026'
}
for idx,text in updates.items():
    p=doc.paragraphs[idx]; p.text=text; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.name='Arial'; r.font.bold=idx in (0,1,4,5); r.font.size=Pt(24 if idx==4 else 11)
        if idx==4: r.font.color.rgb=RGBColor.from_string('17365D')

# Extend the existing static contents list for the consolidated appendices.
toc_last=next(p for p in doc.paragraphs if p.text.strip()=='Appendix B. Test Case Register')
for text in ['Appendix C. CSC2854 Assessment Brief and Compliance Mapping','Appendix D. Complete Test Case Documentation']:
    toc_last.add_run().add_break()
    rr=toc_last.add_run(text); rr.font.name='Arial'; rr.font.size=Pt(11)

# Reuse the exact KPM logo embedded in the supplied assessment PDF.
logo=ROOT/'tmp'/'pdf_template'/'asset-0-Image22.png'
logo_p=doc.paragraphs[0].insert_paragraph_before()
logo_p.alignment=WD_ALIGN_PARAGRAPH.CENTER
logo_p.paragraph_format.space_after=Pt(4)
logo_p.add_run().add_picture(str(logo),width=Inches(1.05))

t=doc.tables[0]
t.cell(1,0).text='Course Name'; t.cell(1,1).text='Project'
t.cell(2,0).text='Course Code'; t.cell(2,1).text='CSC2854'
t.cell(3,0).text='Academic Session'; t.cell(3,1).text='Session 1 2026/2027'
t.cell(4,0).text='Name'; t.cell(4,1).text=''
t.cell(5,0).text='I/D Number / Matric No.'; t.cell(5,1).text=''
t.cell(6,0).text='Class'; t.cell(6,1).text='DCS 6___'
for label,value in [
    ('System','Hotel Maintenance Management System'),
    ('Type of Assessment','Final Project'),
    ('Lecturer','Pn. Rozana Binti Suman / Pn. Nini Aniza Binti Zakaria / Pn. Mawarwiduri Binti Ab Halik')
]:
    cells=t.add_row().cells; cells[0].text=label; cells[1].text=value

for row in t.rows:
    for i,c in enumerate(row.cells):
        for p in c.paragraphs:
            for r in p.runs:
                r.font.name='Arial'; r.font.size=Pt(9); r.bold=(row is t.rows[0] or i==0)

# Add the declaration directly below the student-information block.
anchor=t._tbl
p=OxmlElement('w:p'); pPr=OxmlElement('w:pPr'); jc=OxmlElement('w:jc'); jc.set(qn('w:val'),'center'); pPr.append(jc); p.append(pPr)
r=OxmlElement('w:r'); rPr=OxmlElement('w:rPr'); fonts=OxmlElement('w:rFonts'); fonts.set(qn('w:ascii'),'Arial'); fonts.set(qn('w:hAnsi'),'Arial'); rPr.append(fonts); sz=OxmlElement('w:sz'); sz.set(qn('w:val'),'18'); rPr.append(sz); r.append(rPr)
txt=OxmlElement('w:t'); txt.text='I hereby declare that plagiarism in any form will not be tolerated in this assessment.\n\nSignature: ____________________     Date: ____________________'; r.append(txt); p.append(r)
anchor.addnext(p)

doc.core_properties.title='HMMS Final Report - CSC2854 FYP Template'
doc.core_properties.author=''
doc.save(dst)
print(dst)
